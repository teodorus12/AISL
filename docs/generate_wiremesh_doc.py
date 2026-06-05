"""Word report: Wiremesh (hand wireframe + serial data pipeline)."""

import os
import subprocess
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_DIR = os.path.join(os.path.dirname(__file__), "figures_wiremesh")
OUT_PATH = os.path.join(os.path.dirname(__file__), "Poročilo_wiremesh_AISL.docx")


def ensure_figures():
    if not os.path.isdir(FIG_DIR) or len(os.listdir(FIG_DIR)) < 3:
        subprocess.run([sys.executable, os.path.join(os.path.dirname(__file__), "generate_wiremesh_figures.py")],
                       cwd=ROOT, check=False)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text):
    doc.add_paragraph(text)


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def add_image(doc, name, caption, width=6.0):
    path = os.path.join(FIG_DIR, name)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].italic = True
        doc.add_paragraph()


def build():
    ensure_figures()
    doc = Document()

    t = doc.add_heading("Poročilo: Wiremesh v projektu AISL", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(
        doc,
        "Dokument pokriva dva pomena wire mesha v sistemu AISL: (A) 2D wireframe mrežo "
        "roke pri prepoznavanju znakov in (B) podatkovno mrežo senzorskih chunkov, "
        "ki po serijski povezavi (wire) prihajajo iz STM32 naprave."
    )

    add_heading(doc, "Del A — Wire mesh roke (Hand Tracking)", 1)

    add_heading(doc, "A.1 Kaj je wire mesh roke", 2)
    add_para(
        doc,
        "Wire mesh je mreža črt (robov), ki povezuje 21 landmarkov roke. Uporablja se "
        "za vizualno povratno informacijo uporabniku in ima enako topologijo kot kostni "
        "vektorji v Serializerju (hand_tracking/HT_landmark_drawer.py, HT_serializer.py)."
    )
    add_image(doc, "wm_01_hand_mesh.png", "Slika A1: Topologija wire mesha — indeksi landmarkov")

    add_heading(doc, "A.2 Povezave (21 robov)", 2)
    t = doc.add_table(rows=6, cols=2)
    t.style = "Table Grid"
    rows = [
        ("Prst / del", "Indeksi povezav"),
        ("Palec", "0→1→2→3→4"),
        ("Kazalec", "0→5→6→7→8"),
        ("Sredinec", "5→9→10→11→12"),
        ("Prstanec", "9→13→14→15→16"),
        ("Mehčič", "13→17→18→19→20; 0→17"),
    ]
    for i, (a, b) in enumerate(rows):
        t.rows[i].cells[0].text = a
        t.rows[i].cells[1].text = b
    doc.add_paragraph()

    add_heading(doc, "A.3 Risalni cevovod", 2)
    add_image(doc, "wm_06_render_pipeline.png", "Slika A2: Od kamere do prikaza mesha")
    add_bullet(doc, "LandMarkDrawer: zeleni krogi (landmarki), modre črte (robovi), debelina 2 px.")
    add_bullet(doc, "Koordinate: normalizirane MediaPipe vrednosti × širina/višina slike.")
    add_bullet(doc, "FrameProcessor: horizontalni flip slike (zrcaljenje).")

    add_heading(doc, "A.4 Mesh in AI", 2)
    add_image(doc, "wm_02_mesh_vs_features.png", "Slika A3: Vizualni mesh vs. numerični vektor")

    add_heading(doc, "A.5 Priporočene nadgradnje wire mesha", 2)
    add_bullet(doc, "Uporaba globine z (3D wire mesh) — boljša informacija o orientaciji roke.")
    add_bullet(doc, "Barvno kodiranje prstov za lažjo diagnostiko.")
    add_bullet(doc, "MediaPipe HAND_CONNECTIONS iz uradne knjižnice (enoten vir resnice).")
    add_bullet(doc, "Opcijski 3D prikaz (OpenGL / matplotlib 3D) za poročila.")
    add_bullet(doc, "Prikaz zaupanja detekcije (barva robov ob nizkem confidence).")

    add_heading(doc, "Del B — Wire pipeline (STM32 → podatki)", 1)

    add_heading(doc, "B.1 Pregled", 2)
    add_image(doc, "wm_03_serial_pipeline.png", "Slika B1: Pot podatkov po serijski povezavi")

    add_heading(doc, "B.2 Serijska komunikacija", 2)
    add_bullet(doc, "Vmesnik: UART (pyserial), privzeto COM5, 9600 baud.")
    add_bullet(doc, "Ukazi: GET <datoteka>, STREAM (stream_command.py, get_file.py).")
    add_bullet(doc, "Izhod: binarna datoteka v bin_folder/ ali stream.bin.")
    add_bullet(doc, "Napake: SerialConnectionError, TransferError (errors.py).")

    add_heading(doc, "B.3 Struktura paketov", 2)
    add_image(doc, "wm_04_packet_structure.png", "Slika B2: Shema paketa v BIN dnevniku")
    add_bullet(doc, "Ločilo paketa: 0xFF 0xFF.")
    add_bullet(doc, "Byte stuffing: 0xFE escape sekvence (unstuff_bytes).")
    add_bullet(doc, "CRC16 preverjanje integritete (CRC16.py).")
    add_bullet(doc, "Timestamp v ms, velikost paketa, zaporedje chunkov.")

    add_heading(doc, "B.4 Chunk mesh (senzorski tokovi)", 2)
    add_image(doc, "wm_05_chunk_mesh.png", "Slika B3: Vrste chunkov v paketu")
    t2 = doc.add_table(rows=5, cols=3)
    t2.style = "Table Grid"
    for i, row in enumerate([
        ("Chunk ID", "Vsebina", "Obdelava v Python"),
        ("1", "Pospešek IMU", "int16, 3 osi"),
        ("2", "Žiroskop", "int16, 3 osi"),
        ("3", "Magnetometer", "int16, 3 osi"),
        ("4", "Zvok A-law", "Dekodiranje → WAV (convert_bin_to_wav.py)"),
    ]):
        for j, c in enumerate(row):
            t2.rows[i].cells[j].text = c
    doc.add_paragraph()

    add_heading(doc, "B.5 Povezava z AI moduloma", 2)
    add_bullet(doc, "Chunk 4 (zvok) → WAV → audio_predict.py → beseda.")
    add_bullet(doc, "sign_videos.py: beseda → črke → videi v signs_data/.")
    add_bullet(doc, "Hand tracking: neposredno iz kamere (ne iz STM32).")
    add_bullet(doc, "demo_sis.py: prikaz IMU signalov iz LOG*.BIN (SPO obdelava).")

    add_heading(doc, "B.6 Priporočene nadgradnje wire pipeline", 2)
    add_bullet(doc, "Konfiguracija COM porta (macOS: /dev/cu.*, ne fiksno COM5).")
    add_bullet(doc, "Live STREAM → neposredno v audio AI brez shranjevanja BIN.")
    add_bullet(doc, "Validacija CRC statistika (% izgubljenih paketov).")
    add_bullet(doc, "Sinhronski časovni žig z video hand tracking.")
    add_bullet(doc, "Dokumentacija protokola v ločenem PDF za STM32 firmware.")
    add_bullet(doc, "Web dashboard za prikaz chunk mesh v realnem času.")

    add_heading(doc, "9. Povezane datoteke", 1)
    add_bullet(doc, "hand_tracking/HT_landmark_drawer.py")
    add_bullet(doc, "hand_tracking/HT_serializer.py")
    add_bullet(doc, "create_packets_final.py, stream_command.py, get_file.py")
    add_bullet(doc, "convert_bin_to_wav.py, Signal_decode.py")
    add_bullet(doc, "main.py — meni ukazov 1–11")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("AISL — Wiremesh poročilo").italic = True

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
