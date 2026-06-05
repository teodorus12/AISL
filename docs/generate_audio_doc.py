"""Generate Word document: Prepoznava zvoka v AISL."""

import os
import subprocess
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_DIR = os.path.join(os.path.dirname(__file__), "figures_audio")
OUT_PATH = os.path.join(os.path.dirname(__file__), "Prepoznava_zvoka_AISL.docx")


def ensure_figures():
    if not os.path.isdir(FIG_DIR) or len([f for f in os.listdir(FIG_DIR) if f.endswith(".png")]) < 5:
        subprocess.run(
            [sys.executable, os.path.join(os.path.dirname(__file__), "generate_audio_figures.py")],
            cwd=ROOT,
            check=False,
        )


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def add_numbered(doc, text):
    doc.add_paragraph(text, style="List Number")


def add_image(doc, filename, caption, width=6.0):
    path = os.path.join(FIG_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        doc.add_paragraph()
    else:
        add_para(doc, f"[Slika ni na voljo: {filename}]")


def build():
    ensure_figures()
    doc = Document()

    title = doc.add_heading("Prepoznava zvoka v sistemu AISL", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Tehnična dokumentacija modula za prepoznavo izgovorjenih besed (npr. kava, pivo, čaj) "
        "iz WAV posnetkov. Vključuje analizo algoritmov, grafe in sheme na podlagi "
        "audio_recognition.py, audio_gui.py in povezanih komponent."
    )

    add_heading(doc, "1. Splošni pregled", 1)
    add_para(
        doc,
        "Audio modul pretvori kratki govorni posnetek v vektor akustičnih značilnosti "
        "in z nevronsko mrežo (MLP) določi eno od učenih besed. Rezultat se lahko "
        "poveže z videi znakovnega jezika (črke besede)."
    )
    add_image(doc, "audio_01_pipeline.png", "Slika 1: Potek podatkov — od WAV do besede in videov")

    add_heading(doc, "2. Viri podatkov in datoteke", 1)
    t = doc.add_table(rows=6, cols=2)
    t.style = "Table Grid"
    for i, (a, b) in enumerate([
        ("Mapa / datoteka", "Vloga"),
        ("teaching_data/<beseda>/*.wav", "Učni posnetki za treniranje"),
        ("testing_data/*_test*.wav", "Testni posnetki (ime vsebuje besedo)"),
        ("ai6.pkl", "Shranjen model (privzeto v audio_predict.py)"),
        ("audio_gui.py", "GUI za nalaganje, učenje, test"),
        ("sign_videos.py", "Beseda → črke → predvajanje videov"),
    ]):
        t.rows[i].cells[0].text = a
        t.rows[i].cells[1].text = b
    doc.add_paragraph()

    add_heading(doc, "3. Predobdelava zvoka (AudioProcessor)", 1)
    add_image(doc, "audio_09_preprocessing.png", "Slika 2: Koraki predobdelave")
    add_image(doc, "audio_02_mfcc_spectrogram.png", "Slika 3: Valovna oblika in MFCC spektrogram")

    add_numbered(doc, "Nalaganje: librosa.load(path, sr=8000) — vzorčenje na 8 kHz.")
    add_numbered(doc, "Obrezovanje tišine: trim(y, top_db=20).")
    add_numbered(doc, "Fiksna dolžina: 1 sekunda (8000 vzorcev) — odrez ali zero-padding.")
    add_numbered(doc, "Ekstrakcija značilnosti (librosa).")
    add_numbered(doc, "Agregacija: za vsako značilnost mean in std po času → en vektor.")

    add_heading(doc, "3.1 Ekstrahirane značilnosti", 2)
    add_image(doc, "audio_03_feature_breakdown.png", "Slika 4: Sestava 162-dimenzionalnega vektorja")

    add_bullet(doc, "MFCC (20) + prvi in drugi delta — oblikovanost spektra")
    add_bullet(doc, "Chroma STFT (12) — harmonicna vsebina")
    add_bullet(doc, "Spectral contrast (7) — razlika med vrhovi in dolovami")
    add_bullet(doc, "Zero crossing rate — visokofrekvenčni del")
    add_bullet(doc, "RMS energija — glasnost signala")

    add_para(
        doc,
        "Skupna dimenzija: 2 × (20+20+20+12+7+1+1) = 162. Ob učenju se uporabi "
        "Z-score normalizacija (mean, std) shranjena v modelu."
    )

    add_heading(doc, "3.2 Analiza algoritma", 2)
    add_para(
        doc,
        "Pristop je klasična pipeline za govorovno prepoznavo na kratkih izrazih: "
        "ne uporablja globokih CNN nad spektrogramom, temveč ročno izbrane deskriptorje. "
        "Prednosti: hitrost, majhen model, razumljivost. Omejitve: občutljivost na "
        "šum, mikrofon, naglas in variacijo librosa različice."
    )

    add_heading(doc, "4. Nevronska mreža (NeuralNetwork)", 1)
    add_image(doc, "audio_04_nn_architecture.png", "Slika 5: Arhitektura MLP")

    add_heading(doc, "4.1 Forward pass", 2)
    add_bullet(doc, "Normalizacija: X_norm = (X − mean) / std")
    add_bullet(doc, "Skriti sloj 1: ReLU(X @ W1 + b1)")
    add_bullet(doc, "Skriti sloj 2: ReLU(a1 @ W2 + b2)")
    add_bullet(doc, "Izhod: softmax(a2 @ W3 + b3) → verjetnosti po razredih")

    add_heading(doc, "4.2 Učenje", 2)
    add_bullet(doc, "Izguba: križna entropija (cross-entropy)")
    add_bullet(doc, "Dropout 0,3 v skritih slojih (samo pri učenju)")
    add_bullet(doc, "Gradientni spust po mini-batch (privzeto 32)")
    add_bullet(doc, "Learning rate 0,003 z decay 0,995 na epoho")
    add_bullet(doc, "Privzeto 200 epoh (audio_gui.py)")

    add_heading(doc, "4.3 Privzeti hiperparametri (GUI)", 2)
    t2 = doc.add_table(rows=7, cols=2)
    t2.style = "Table Grid"
    for i, (a, b) in enumerate([
        ("Parameter", "Privzeta vrednost"),
        ("Hidden layer 1", "128 nevronov"),
        ("Hidden layer 2", "64 nevronov"),
        ("Learning rate", "0,003"),
        ("Dropout", "0,3"),
        ("Epochs", "200"),
        ("Batch size", "32"),
    ]):
        t2.rows[i].cells[0].text = a
        t2.rows[i].cells[1].text = b
    doc.add_paragraph()

    add_heading(doc, "5. Učna množica", 1)
    add_image(doc, "audio_05_dataset.png", "Slika 6: Porazdelitev posnetkov v teaching_data/")
    add_para(
        doc,
        "Učni razredi v projektu: kava, pivo, sok, vino, čaj. Skupaj približno "
        "1600 posnetkov različnih govorcov (matej, teo, miha, …). "
        "Za uravnotežen model je priporočljivo približno enako število posnetkov na besedo."
    )

    add_heading(doc, "6. Evalvacija in natančnost", 1)
    add_image(doc, "audio_06_training_curves.png", "Slika 7: Loss in natančnost med učenjem (ilustrativno)")
    add_image(doc, "audio_07_confusion_matrix.png", "Slika 8: Matrika zamenjav (test 20 %, teaching_data)", width=5.5)
    add_image(doc, "audio_08_f1_scores.png", "Slika 9: F1-score po besedah")

    add_heading(doc, "6.1 Metrike", 2)
    add_bullet(doc, "Precision — delež pravilnih napovedi med vsemi napovedmi dane besede")
    add_bullet(doc, "Recall — delež ujetih primerov dejanske besede")
    add_bullet(doc, "F1-score — harmonično povprečje obeh")
    add_para(
        doc,
        "Grafe 8–9 generira skripta z delitvijo teaching_data (80/20) in kratkim "
        "ponovnim učenjem z enako arhitekturo. Shranjen model ai6.pkl lahko doseže "
        "drugačne rezultate — za uradni test uporabite tests_for_ai.py na testing_data/."
    )

    add_heading(doc, "6.2 Testiranje", 2)
    add_para(doc, "Ukazi:", bold=True)
    add_bullet(doc, "python audio_gui.py — GUI: Load Training Data → Train → Test WAV")
    add_bullet(doc, "python tests_for_ai.py — batch test na testing_data/ z ai6.pkl")
    add_bullet(doc, "python main.py → možnost 10 — prepoznava + predvajanje znakov")
    add_bullet(doc, "python audio_predict.py <pot_do.wav> — napoved ene datoteke")

    add_heading(doc, "7. Integracija z znakovnim jezikom", 1)
    add_image(doc, "audio_10_sign_videos.png", "Slika 10: Od besede do videov po črkah")

    add_heading(doc, "8. Algoritmična analiza", 1)
    add_image(doc, "audio_11_complexity.png", "Slika 11: Časovna zahtevnost po fazah")

    add_heading(doc, "9. Primerjava z hand tracking modulom", 1)
    t3 = doc.add_table(rows=6, cols=3)
    t3.style = "Table Grid"
    for i, row in enumerate([
        ("", "Audio", "Hand tracking"),
        ("Vhod", "Zvok (WAV)", "Video (kamera)"),
        ("Značilnosti", "MFCC, chroma, … (162)", "Landmarki roke (126)"),
        ("Model", "Lastni MLP (NumPy)", "sklearn MLP"),
        ("Razredi", "Besede (5+)", "Črke A–Z"),
        ("Izhod", "Beseda + videi črk", "Črka + zaupanje"),
    ]):
        for j, cell in enumerate(row):
            t3.rows[i].cells[j].text = cell
    doc.add_paragraph()

    add_heading(doc, "10. Omejitve in priporočila", 1)
    add_bullet(doc, "Posnemajte v tihem okolju, enako razdaljo do mikrofona.")
    add_bullet(doc, "Vsaj 50–100 posnetkov na besedo za stabilen model.")
    add_bullet(doc, "Po dodajanju besed ponovno trenirajte in shranite nov .pkl.")
    add_bullet(doc, "Preverite združljivost librosa (spectral_contrast pri 8 kHz).")
    add_bullet(doc, "Za produkcijo razmislite o ločenem validacijskem setu (testing_data).")

    add_heading(doc, "11. Datoteke v projektu", 1)
    add_bullet(doc, "audio_recognition.py — NeuralNetwork + AudioProcessor")
    add_bullet(doc, "audio_gui.py — grafično učenje in test")
    add_bullet(doc, "audio_predict.py — napoved iz WAV")
    add_bullet(doc, "tests_for_ai.py — evalvacija na testing_data")
    add_bullet(doc, "sign_videos.py — povezava z videi znakov")
    add_bullet(doc, "docs/figures_audio/ — slike v tem dokumentu")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Projekt AISL — Audio modul | Dokument generiran iz kode in teaching_data").italic = True

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
