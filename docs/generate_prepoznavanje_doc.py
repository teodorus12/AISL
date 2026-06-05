"""Generate Word document with analyses, graphs and diagrams."""

import os
import subprocess
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_DIR = os.path.join(os.path.dirname(__file__), "figures")
OUT_PATH = os.path.join(os.path.dirname(__file__), "Prepoznavanje_znakov_hand_tracking.docx")


def ensure_figures():
    if not os.path.isdir(FIG_DIR) or len(os.listdir(FIG_DIR)) < 5:
        subprocess.run([sys.executable, os.path.join(os.path.dirname(__file__), "generate_figures.py")],
                       cwd=ROOT, check=False)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def add_numbered(doc, text):
    doc.add_paragraph(text, style="List Number")


def add_image(doc, filename, caption, width=6.0):
    path = os.path.join(FIG_DIR, filename)
    if not os.path.exists(path):
        # learning curve may be in models/
        alt = os.path.join(ROOT, "models", filename)
        path = alt if os.path.exists(alt) else path
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        doc.add_paragraph()
    else:
        add_para(doc, f"[Slika ni na voljo: {filename}]")


def build():
    ensure_figures()
    doc = Document()

    title = doc.add_heading("Prepoznavanje znakov v sistemu AISL Hand Tracking", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Tehnična dokumentacija z opisom algoritmov, analizo natančnosti, "
        "shemit poteka in grafi na podlagi dejanske implementacije ter učne množice projekta."
    )

    # 1
    add_heading(doc, "1. Splošni pregled", 1)
    add_para(
        doc,
        "Sistem prepoznava statistične črke slovenskega znakovnega jezika iz oblike roke "
        "v posameznem video okviru. Ne uporablja surovih pikslov, temveč 21 točk MediaPipe "
        "in 126-dimenzionalni vektor geometrije."
    )
    add_image(doc, "01_pipeline.png", "Slika 1: Arhitektura sistema — potek podatkov od kamere do napovedi")

    # 2 MediaPipe
    add_heading(doc, "2. Zaznava roke — MediaPipe Hand Landmarker", 1)
    add_image(doc, "02_hand_skeleton.png", "Slika 2: 21 landmarkov in povezave (enake kot v aplikaciji)")

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    for i, (a, b) in enumerate([
        ("Komponenta", "Opis"),
        ("Vhod", "Slika iz kamere (BGR → RGB)"),
        ("Izhod", "21 landmarkov (x, y, z)"),
        ("Model", "hand_landmarker.task (TFLite, prednaučen)"),
        ("Nastavitve", "1 roka; min. confidence 0,5"),
    ]):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
    doc.add_paragraph()

    add_heading(doc, "2.1 Analiza algoritma detekcije", 2)
    add_bullet(doc, "Detekcija in sledenje sta ločena: ob izgubi roke sledenje ponastavi detekcijo.")
    add_bullet(doc, "Koordinate so normalizirane na dimenzije slike (0–1), ne v pikslih.")
    add_bullet(doc, "Napaka v eni točki vpliva na celoten feature_vector in posledično na MLP.")

    # 3 Serializer
    add_heading(doc, "3. Pretvorba v značilnosti (Serializer)", 1)
    add_image(doc, "05_normalization_flow.png", "Slika 3: Zaporedje transformacij v Serializerju")
    add_image(doc, "03_feature_breakdown.png", "Slika 4: Sestava 126-dimenzionalnega vektorja")

    add_heading(doc, "3.1 Matematični opis normalizacije", 2)
    add_para(doc, "Naj bodo p_i = (x_i, y_i, z_i) koordinate landmarka i, p_0 zapestje:")
    add_bullet(doc, "Relativne koordinate: p'_i = p_i − p_0")
    add_bullet(doc, "Skala: s = ||p'_9|| (osnova srednjega prsta); p''_i = p'_i / s")
    add_bullet(doc, "Kostni vektor med točkama a in b: v = (p''_b − p''_a) / ||p''_b − p''_a||")
    add_bullet(doc, "feature_vector = flatten(p'') ⊕ flatten(v) ∈ R^126")

    add_heading(doc, "3.2 Zakaj ta pristop deluje", 2)
    add_bullet(doc, "Translacijska invarianca: premik roke v kadru ne spremeni p''.")
    add_bullet(doc, "Skalna invarianca: približanje/oddalitev kamere delno kompenzira s.")
    add_bullet(doc, "Omejitev: brez rotacijske invariacije — drugačen kot roke lahko zmanjša natančnost.")

    # 4 Training
    add_heading(doc, "4. Učna množica in treniranje", 1)
    add_image(doc, "06_dataset_distribution.png", "Slika 5: Porazdelitev posnetkov po razredih (projekt AISL)")
    add_image(doc, "04_mlp_architecture.png", "Slika 6: Arhitektura nevronske mreže MLP")

    t2 = doc.add_table(rows=7, cols=2)
    t2.style = "Table Grid"
    for i, (a, b) in enumerate([
        ("Korak", "Metoda"),
        ("Agregacija posnetka", "Povprečje (mean) vseh feature_vector v JSON posnetku"),
        ("Model", "MLPClassifier: 256 → 128 → 64, ReLU, softmax"),
        ("Optimizacija", "Adam (privzeto v sklearn), max 500 iteracij"),
        ("Delitev", "80 % train / 20 % test, stratified"),
        ("Metrike", "precision, recall, F1, confusion matrix"),
        ("Izhod", "models/sl_model.pkl, models/sl_classes.json"),
    ]):
        t2.rows[i].cells[0].text = a
        t2.rows[i].cells[1].text = b
    doc.add_paragraph()

    add_heading(doc, "4.1 Analiza učne množice", 2)
    add_para(
        doc,
        "Graf na sliki 5 prikazuje dejansko število posnetkov na črko v mapi Handtracking/clips/. "
        "Neenakomerna porazdelitev (npr. več posnetkov za I kot za druge črke) lahko povzroči, "
        "da je model boljši pri pogostih razredih. Priporočilo: vsaj 15–20 posnetkov na črko."
    )

    add_image(doc, "09_learning_curve.png", "Slika 7: Učna krivulja (training loss) po treniranju", width=5.5)

    # 5 Evaluation
    add_heading(doc, "5. Evalvacija natančnosti modela", 1)
    add_para(
        doc,
        "Spodnji grafi temeljijo na isti metodi kot HT_train_model.py: 20 % posnetkov "
        "je rezerviranih za test, model pa ni bil učen na njih."
    )
    add_image(doc, "07_confusion_matrix.png", "Slika 8: Matrika zamenjav na testni množici", width=6.2)
    add_image(doc, "08_f1_per_class.png", "Slika 9: F1-score po posameznem razredu")

    add_heading(doc, "5.1 Interpretacija rezultatov", 2)
    add_bullet(doc, "Diagonalna matrika (slika 8): pravilne napovedi; izven diagonale: zamenjave.")
    add_bullet(doc, "F1 < 0,7 (oranžno/rdeče): razred potrebuje več ali bolj raznolikih posnetkov.")
    add_bullet(doc, "Visok F1 na testu ne garantira enake uspešnosti v live načinu (glej poglavje 6).")

    # 6 Live
    add_heading(doc, "6. Prepoznavanje v realnem času", 1)
    add_image(doc, "11_train_vs_live.png", "Slika 10: Razlika med učenjem (mean) in live (posamezen okvir)")
    add_image(doc, "10_live_smoothing.png", "Slika 11: Primer glajenja z majority vote (10 okvirjev)")

    add_numbered(doc, "Izračun feature_vector za trenutni okvir (~100 Hz).")
    add_numbered(doc, "MLP.predict_proba → verjetnosti po razredih.")
    add_numbered(doc, "Če max(proba) < 0,6 → prikaz »?«.")
    add_numbered(doc, "Sicer majority vote nad zadnjimi 10 napovedmi.")
    add_numbered(doc, "Brez roke: reset zgodovine, prikaz »—«.")

    # 7 Complexity
    add_heading(doc, "7. Algoritmična analiza", 1)
    add_image(doc, "12_complexity.png", "Slika 12: Povzetek časovne in prostorske zahtevnosti")

    add_heading(doc, "7.1 Ocena zmogljivosti v praksi", 2)
    add_bullet(doc, "MediaPipe + MLP na CPU: tipično 30–60 FPS (odvisno od stroja).")
    add_bullet(doc, "Ozko grlo: običajno detekcija MediaPipe, ne MLP.")
    add_bullet(doc, "Treniranje: sekunde do minut za ~700 posnetkov.")

    # 8 Limitations
    add_heading(doc, "8. Omejitve in priporočila", 1)
    add_bullet(doc, "Samo statična drža — ne dinamika znaka skozi čas.")
    add_bullet(doc, "Ena roka na kadar.")
    add_bullet(doc, "Live uporablja okvir, učenje povprečje posnetka — držite znak mirno pri snemanju.")
    add_bullet(doc, "Po novih posnetkih: python hand_tracking/HT_train_model.py in ponovni zagon aplikacije.")

    # 10 Upgrades
    add_heading(doc, "10. Priporočene nadgradnje in dopolnitve", 1)
    add_para(
        doc,
        "Spodnji predlogi temeljijo na trenutni arhitekturi (statistična črka iz ene "
        "drže roke) in na opaženih omejitvah v kodi ter pri uporabi."
    )
    add_image(doc, "13_upgrade_roadmap.png", "Slika 13: Predlagan načrt nadgradnje (časovnica)")

    add_heading(doc, "10.1 Visoka prioriteta (najprej)", 2)
    t_up = doc.add_table(rows=6, cols=3)
    t_up.style = "Table Grid"
    for i, row in enumerate([
        ("Področje", "Kaj dodati / popraviti", "Zakaj"),
        ("Učenje vs. live", "Enaka agregacija: ali mean več okvirjev v live, ali učenje na posameznih okvirjih", "Zmanjša napake pri premikanju roke"),
        ("Podatki", "15–30 posnetkov na črko, več kotov in razdalj", "Višji F1, manj zamenjav"),
        ("Model", "Ponovno treniranje po vsaki novi črki; preveriti sl_classes.json", "Nov znak se ne pojavi brez train + restart"),
        ("GUI", "Fiksni desni panel (že izvedeno v HT_window.py)", "Uporabnik vidi AI napoved"),
        ("Koda", "Enoten izvor povezav mesh (MediaPipe HAND_CONNECTIONS)", "Serializer in LandMarkDrawer ne divergirata"),
    ]):
        for j, c in enumerate(row):
            t_up.rows[i].cells[j].text = c
    doc.add_paragraph()

    add_heading(doc, "10.2 Srednja prioriteta — algoritmi", 2)
    add_bullet(doc, "Časovni model: LSTM, GRU ali 1D-CNN nad zaporedjem feature_vector (namesto samo mean posnetka).")
    add_bullet(doc, "Augmentacija: majhen šum na landmarkih, rotacija v ravnini, simulacija pomanjkanja točk.")
    add_bullet(doc, "Ločen validacijski set (mapa val/) — ne ocenjevati samo na training accuracy.")
    add_bullet(doc, "Class weights v MLP pri neenakomernih razredih (nekaterih črk je več posnetkov).")
    add_bullet(doc, "Kalibracija praga zaupanja (60 %) glede na validation ROC krivuljo.")
    add_bullet(doc, "Shranjevanje napačnih napovedi za ponovno označevanje (active learning).")

    add_heading(doc, "10.3 Srednja prioriteta — funkcionalnost", 2)
    add_bullet(doc, "Podpora za obe roki (num_hands=2) z ločenim labelom (Leva/Desna + črka).")
    add_bullet(doc, "Prepoznava besed kot zaporedje črk (povezava s sign_videos.py).")
    add_bullet(doc, "Izvoz metrik v CSV/JSON ob treniranju (ne samo classification_report v terminalu).")
    add_bullet(doc, "Način »demo« brez kamere — predvajanje iz shranjenih JSON posnetkov.")

    add_heading(doc, "10.4 Dolgoročno — raziskovalne smeri", 2)
    add_bullet(doc, "Dinamični znaki SZJ: DTW, Temporal CNN ali Transformer nad sekvenco okvirjev.")
    add_bullet(doc, "3D wire mesh z globino (z) in boljša rotacijska invarianca.")
    add_bullet(doc, "MediaPipe Holistic (telo + obraz) za kontekst pri veččlennih znakih.")
    add_bullet(doc, "Edge deployment (TFLite / ONNX) za hitrejši inference brez težkega MediaPipe strežnika.")
    add_bullet(doc, "Integracija z audio + STM32: sinhron zvok–video za večmodalno napoved.")

    add_heading(doc, "10.5 Tehnični dolg v projektu", 2)
    add_bullet(doc, "Duplikat map Handtracking/clips in hand_tracking/Handtracking/clips — enotna pot.")
    add_bullet(doc, "Model path: models/ vs hand_tracking/models/ — uskladiti privzete poti v HT_ai.py.")
    add_bullet(doc, "requirements.txt z verzijami (opencv, mediapipe, sklearn, librosa).")
    add_bullet(doc, "Avtomatski test: test_hand_tracking.py z mock landmarki.")

    add_heading(doc, "10.6 Merila uspeha po nadgradnji", 2)
    add_bullet(doc, "F1 > 0,9 na validacijskem setu za vse razrede.")
    add_bullet(doc, "Live napoved stabilna ≥ 2 s brez utripanja (?) pri mirni roki.")
    add_bullet(doc, "Čas odprtja aplikacije < 3 s; inference > 25 FPS na ciljni strojni opremi.")

    # 11 Files
    add_heading(doc, "11. Datoteke v projektu", 1)
    add_bullet(doc, "hand_tracking/HT_handler.py — MediaPipe")
    add_bullet(doc, "hand_tracking/HT_serializer.py — vektorizacija")
    add_bullet(doc, "hand_tracking/HT_train_model.py — učenje")
    add_bullet(doc, "hand_tracking/HT_ai.py — live prepoznavanje")
    add_bullet(doc, "Handtracking/clips/ — učni posnetki")
    add_bullet(doc, "models/sl_model.pkl — naučen model")
    add_bullet(doc, "docs/figures/ — slike v tem dokumentu")
    add_bullet(doc, "docs/Poročilo_wiremesh_AISL.docx — wire mesh in wire pipeline")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Projekt AISL — Hand Tracking | Dokument generiran iz kode in učne množice").italic = True

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
